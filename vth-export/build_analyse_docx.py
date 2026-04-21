"""Genereer een nette .docx van de VTH OmgevingsCheck strategische analyse."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLAUW = RGBColor(0x1B, 0x4B, 0x6B)
DONKERGRIJS = RGBColor(0x33, 0x33, 0x33)
ACCENT = RGBColor(0xE6, 0x51, 0x00)

doc = Document()

# --- standaard stijl ---
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = DONKERGRIJS

# --- pagina-marges ---
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLAUW
        run.font.name = 'Calibri'
    return h

def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    return p

# ===== TITELPAGINA =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('VTH OmgevingsCheck')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = BLAUW

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Strategische App-analyse')
run.font.size = Pt(18)
run.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Toepassingen · Kansen · Valkuilen · Toekomstige mogelijkheden')
run.italic = True
run.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Datum: 21 april 2026')
run.font.size = Pt(11)

doc.add_page_break()

# ===== 1. Wat is deze app in één zin =====
heading('1. Wat is deze app in één zin?', level=1)
para(
    'Een offline-first single-file tabletapp voor gemeentelijk bouwtoezicht volgens het Besluit '
    'bouwwerken leefomgeving (Bbl), waarmee een toezichthouder zaken beheert, controles op locatie '
    'uitvoert, foto’s koppelt aan bouwonderwerpen en getoetst wordt tegen 79 geverifieerde '
    'Bbl-artikelen en 28 toetsniveaus uit het BWTinfo-toezichtprotocol.'
)
para('Schaal: circa 10.200 regels HTML/CSS/JS in één bestand, ~600 KB, geen build-stap, geen backend.', italic=True)

# ===== 2. Functieoverzicht =====
heading('2. Volledig functieoverzicht (huidige toepassingen)', level=1)

heading('2.1 Zaak- en projectbeheer', level=2)
for t in [
    'Zakenlijst met filterchips (type zaak + status) en badges met counts',
    'Sorteren op laatste controle, aanmaakdatum, naam of status',
    'Zoeken op naam, adres of zaaknummer',
    'Splitview-detail (lijst + detail) met 3 tabs: Basisgegevens, Controles, Voortgang',
    'Kaartweergave (Leaflet) met status-gekleurde markers',
    '"Vandaag"-zijpaneel met dagagenda',
]:
    bullet(t)

heading('2.2 Controles op locatie', level=2)
for t in [
    '6 controle-soorten: aanloop-, tussentijdse, eind-, melding-, her- en handhavingscontrole',
    'Onderwerp-gedreven toetsing met resultaat (akkoord / akkoord-mits / niet-akkoord)',
    'Bbl-artikel-lookup: 79 geverifieerde artikelen met lid-structuur en volledige wettekst',
    'Toetsniveau-koppeling: automatische mapping naar BWTinfo-toetsmomenten',
    'Vrije tekstvelden voor afwijking, termijn en advies aan aannemer/handhaving',
    'Startuitvoering-detectie zet projectstatus automatisch op "In uitvoering"',
]:
    bullet(t)

heading('2.3 Foto’s', level=2)
for t in [
    'Per controle én per onderwerp koppelbaar',
    'Base64-opslag in localStorage (vth_fotos)',
    'Toelichting per foto',
    'Native camera-toegang via capture="environment"',
]:
    bullet(t)

heading('2.4 PDF-import (vergunningen)', level=2)
for t in [
    'Upload vergunning-PDF → tekstextractie via PDF.js (offline worker)',
    'Gebruiker selecteert tekst via muis/touch-sleep en wijst toe aan velden (zaaknummer, adres, voorschriften, etc.)',
    'Importdata vult het projectformulier vooringevuld in',
]:
    bullet(t)

heading('2.5 Externe koppelingen', level=2)
for t in [
    'Ruimtelijke Plannen (DSO) — zoekopdracht op RD- of ETRS89-coördinaten',
    'PDOK geocoding — automatisch adres → lat/lng + RD bij XLS-import',
    'Zaaksysteem-knop — externe link (placeholder, uitbreidbaar)',
]:
    bullet(t)

heading('2.6 Rapportage en export', level=2)
for t in [
    'Print-view met gemeentenaam en logo, html2pdf-export',
    'Sjabloonbeheer: eigen controlesjablonen definiëren',
    'JSON-backup (manueel + automatisch) van projecten + foto’s + settings',
    'XLS-import van zakenexport met PDOK-geocoding op de achtergrond',
]:
    bullet(t)

heading('2.7 Instellingen', level=2)
for t in [
    'Gemeentenaam, standaard toezichthouder, logo',
    'API-key voor Ruimtelijke Plannen (DSO)',
]:
    bullet(t)

# ===== 3. Kansen =====
heading('3. Kansen (Opportunities)', level=1)

kansen = [
    ('Nederlandse gemeentemarkt', 'Circa 340 gemeenten, allen gebonden aan de Omgevingswet. Geen SaaS-concurrent is volledig offline + tabletgericht.'),
    ('Offline-first als verkoopargument', 'Kelders, bouwputten en kraanplaatsen hebben vaak geen netwerk; concurrenten (Squit, Centric, Roxit) zijn online-afhankelijk.'),
    ('Bbl-artikelen als aanjager', '79 artikelen is een uniek inhoudelijk actief. Uitbreiden tot volledige hoofdstukken 2–4 maakt dit "de Bbl-bijbel voor inspecteurs".'),
    ('Toetsprotocol-integratie BWTinfo', 'Sluit direct aan op landelijk geaccepteerde werkwijze.'),
    ('DSO/Ruimtelijke Plannen koppeling', 'Kan uitgebouwd tot automatische strijdigheidscheck tussen omgevingsplan en controle.'),
    ('Wetgevingsbestanden als kennisbasis', 'De aanwezige BWBR0041297-tekstbestanden (hoofdstuk 3 en 4) vormen basis voor zoekmachine of AI-RAG.'),
    ('Platformisering', 'Het datamodel is generiek genoeg om milieu, APV, sloop en brandveiligheid te dekken.'),
    ('Laagdrempelige uitrol', 'Geen server, geen IT-afstemming. Pilot bij een gemeente kan binnen een week draaien.'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Kans'
hdr[1].text = 'Waarom dit telt'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_shading(cell, '1B4B6B')
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for titel, toelichting in kansen:
    row = table.add_row().cells
    row[0].text = titel
    row[1].text = toelichting
    for p in row[0].paragraphs:
        for r in p.runs:
            r.bold = True

# ===== 4. Valkuilen =====
heading('4. Valkuilen (risico’s en technische schuld)', level=1)

heading('4.1 Data en privacy', level=2)
for t in [
    'Geen authenticatie — iedereen met tablettoegang ziet alle zaken',
    'Base64-foto’s in localStorage — browserlimiet ±5-10 MB, bij ~50 foto’s per project raakt opslag vol en faalt setItem stil',
    'AVG-risico: persoonsgegevens (initiatiefnemer, aannemer, foto’s van personen) staan onversleuteld in browser-storage',
    'Data verdwijnt bij cache-clear of browserwissel — geen echte sync, alleen handmatige JSON-export',
    'Geen audit log — wie heeft wanneer welk oordeel gegeven is niet herleidbaar',
]:
    bullet(t)

heading('4.2 Technische schuld', level=2)
for t in [
    '10.200 regels in één bestand, geen modules, geen tests, geen build-tool',
    'Duplicatie tussen index.html (536 KB) en index-tablet.html (600 KB)',
    '15+ globale state-variabelen en monkey-patched navigatie (_origShowScreen)',
    'Legacy-datamigratie was nodig — bewijst dat datacontract-evolutie ongecontroleerd verloopt',
    'Inconsistente storage-keys (vth_projecten vs vth_projects)',
    'DATAMODEL.md is verouderd en beschrijft nog oude status-/typewaarden',
    'IndexedDB-sync en FileSystemFileHandle zijn half geïmplementeerd (dode code)',
]:
    bullet(t)

heading('4.3 Functionele kwetsbaarheden', level=2)
for t in [
    'PDF-tekstextractie met Y-positie-threshold 3px — breekt bij complexe vergunningen (tabellen, 2-koloms)',
    'Geen duplicatendetectie bij XLS-import',
    'Geen validatie van verplichte velden buiten "naam" en "type zaak"',
    'Rapportgeneratie via html2pdf is traag bij veel foto’s en kan vastlopen',
]:
    bullet(t)

heading('4.4 Organisatorisch en commercieel', level=2)
for t in [
    'One-person-codebase zonder documentatie van architectuurbeslissingen',
    'Geen versiebeheer voor Bbl-artikelen — jaarlijkse wetswijzigingen vereisen handmatige update',
    'Geen SLA, supportproces of release-cyclus gedefinieerd',
    'Vendor lock-in op PDF.js en html2pdf (html2pdf is CDN-afhankelijk)',
]:
    bullet(t)

# ===== 5. Toekomstige toepassingen =====
heading('5. Toekomstige toepassingen — uitgewerkte roadmap', level=1)

heading('Horizon 1 — Stabilisatie (1-3 maanden)', level=2)
h1 = [
    'Cloud-backup via Azure/AWS S3 met gemeente-gebonden bucket — offline blijft primair',
    'IndexedDB-migratie voor foto’s — blob-opslag i.p.v. base64, 100× meer ruimte',
    'Automatische schema-migraties als herbruikbaar framework',
    'index.html en index-tablet.html samenvoegen of mobiele variant definitief verwijderen',
    'Rollenmodel: toezichthouder / teamleider / gemeente-admin met simpele PIN-login',
    'Wetsversie-flag: Bbl-artikelen krijgen geldigVanaf/geldigTot velden',
]
for i, t in enumerate(h1, 1):
    numbered(t)

heading('Horizon 2 — Gebruikswaarde vergroten (3-9 maanden)', level=2)
h2 = [
    'AI-assist bij PDF-import — LLM extraheert voorschriften automatisch',
    'AI-samenvatting van controlebevindingen — dicteren, AI structureert naar advies/afwijking/termijn',
    'QR-code op bouwbord — scannen opent direct de juiste zaak op locatie',
    'Spraakmemo’s tijdens controle met Whisper-transcriptie (offline mogelijk)',
    'Beeldvergelijking voor/na per onderwerp — side-by-side bij hercontrole',
    'Automatische strijdigheidscheck omgevingsplan via DSO-API bij openen project',
    'Offline kaarttiles via pre-cache van gemeentegrens (PDOK BRT)',
    'Digitale handtekening op controlerapport (toezichthouder + aannemer)',
    'Handhavingsbrief-generator uit "niet-akkoord"-bevindingen',
    'Herinneringen voor termijngebonden hercontroles via Web Push',
]
for t in h2:
    numbered(t)

heading('Horizon 3 — Platformisering (9-18 maanden)', level=2)
h3 = [
    'Multi-domein modules: milieu, APV, sloop, brandveiligheid',
    'Koppeling met zaaksystemen: RX Mission, Djuma, JOIN, Centric iVergunningen',
    'Stakeholder-portal voor aannemer: eigen afwijkingen inzien, reactie uploaden',
    'Management-dashboard: doorlooptijden, compliance-rate, handhavings-incidenten per wijk',
    'Native app-wrappers (Capacitor of WKWebView) voor betere camera- en GPS-integratie',
    'Koppeling met LVBB (Landelijke Voorziening Bekendmaken en Beschikbaar stellen)',
    'Uitbreiding BBL_ARTIKELEN naar volledige hoofdstukken 2-7',
    'AI-kennisbank-chat (RAG op aanwezige BWBR-bestanden)',
]
for t in h3:
    numbered(t)

heading('Horizon 4 — Innovatie (18+ maanden)', level=2)
h4 = [
    'Computer vision op bouwfoto’s — automatische detectie van veiligheidsrisico’s',
    'Drone-integratie: inspectie-foto’s direct uit drone-missie koppelen aan onderwerpen',
    'IoT-sensoren: temperatuur/vocht bij betonstort, trilling bij heien',
    'Wetsveranderingen-monitor: automatische flag bij wijziging van geraadpleegde artikelen',
    'Landelijke geanonimiseerde benchmarking tussen gemeenten',
]
for t in h4:
    numbered(t)

# ===== 6. Prioriteitsadvies =====
heading('6. Prioriteitsadvies', level=1)

p = doc.add_paragraph()
run = p.add_run('Als er maar één vervolgstap gekozen mag worden: ')
run.bold = True
p.add_run(
    'IndexedDB-migratie voor foto’s (Horizon 1, punt 2). Dit is de enige valkuil die '
    'vandaag al leidt tot stil dataverlies bij een toezichthouder die meer dan 20 foto’s per dag maakt.'
)

p = doc.add_paragraph()
run = p.add_run('Daarna: ')
run.bold = True
p.add_run(
    'Rollenmodel en cloud-backup (Horizon 1, punten 1 en 5). Zonder deze is de app juridisch '
    'niet houdbaar bij een AVG-audit.'
)

p = doc.add_paragraph()
run = p.add_run('Hoogste business-waarde bij gelijke inspanning: ')
run.bold = True
p.add_run(
    'AI-voorschriftextractie en automatische strijdigheidscheck DSO (Horizon 2, punten 1 en 6). '
    'Deze features onderscheiden de app van elk bestaand vergunningssysteem in Nederland.'
)

# ===== VOETREGEL =====
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run('VTH OmgevingsCheck · Strategische App-analyse · april 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

out_path = '/home/user/demo/vth-export/VTH_OmgevingsCheck_Strategische_Analyse.docx'
doc.save(out_path)
print(f'Document opgeslagen: {out_path}')
